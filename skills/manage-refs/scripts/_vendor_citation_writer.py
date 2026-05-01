# SPDX-License-Identifier: MIT
# Copyright (c) 2026 Ali Soroush
# Vendored from https://github.com/alisoroushmd/zotero-mcp
#   src/zotero_mcp/citation_writer.py @ ed5dfb71
# Imported into medsci-skills 2026-05-01 (originally vendored into 01_RFA_Adjunct
# the same day, relocated here after RFA-Adjunct validation). No functional
# modifications — `inject_zotero_cwyw.py` patches `zotero_to_csl_json` at import
# time to use Zotero's native `?format=csljson` endpoint, which handles
# webpage / report / non-journal item types correctly. See ../NOTICE.md.
# Full license: ../LICENSE.zotero-mcp

"""Citation writer -- builds Word documents with live Zotero field codes.

Creates .docx files containing ADDIN ZOTERO_ITEM and ADDIN ZOTERO_BIBL
field codes that the Zotero Word plugin recognizes as live citations.
"""

from __future__ import annotations

import hashlib
import json
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# 1. Citation Parser
# ---------------------------------------------------------------------------

_CITATION_RE = re.compile(r"\[@([A-Za-z0-9]+(?:,\s*@[A-Za-z0-9]+)*)\]")


@dataclass
class TextBlock:
    """A segment of parsed text -- either plain text or a citation group."""

    kind: str  # "text" or "citation"
    content: str  # raw text for "text", empty for "citation"
    keys: list[str] = field(default_factory=list)
    numbers: list[int] = field(default_factory=list)


def parse_citations(text: str) -> tuple[list[TextBlock], dict[str, int]]:
    """Parse text with [@KEY] markers into blocks with Vancouver numbering.

    Supports single ``[@KEY]`` and grouped ``[@KEY1, @KEY2]`` citations.
    Numbers are assigned sequentially by order of first appearance.
    Same key reuses its number.

    Args:
        text: Input text containing citation markers.

    Returns:
        Tuple of (list of TextBlocks, dict mapping item_key -> vancouver_number).
    """
    blocks: list[TextBlock] = []
    mapping: dict[str, int] = {}
    counter = 0
    last_end = 0

    for match in _CITATION_RE.finditer(text):
        start, end = match.span()

        # Text before this citation
        if start > last_end:
            blocks.append(TextBlock("text", text[last_end:start]))

        # Parse keys from the match group
        raw_keys = match.group(1)
        keys = [k.strip().lstrip("@") for k in raw_keys.split(",")]

        numbers: list[int] = []
        for key in keys:
            if key not in mapping:
                counter += 1
                mapping[key] = counter
            numbers.append(mapping[key])

        blocks.append(TextBlock("citation", "", keys, numbers))
        last_end = end

    # Trailing text
    if last_end < len(text):
        blocks.append(TextBlock("text", text[last_end:]))

    return blocks, mapping


# ---------------------------------------------------------------------------
# 2. CSL-JSON Converter
# ---------------------------------------------------------------------------

_ITEM_TYPE_MAP: dict[str, str] = {
    "journalArticle": "article-journal",
    "book": "book",
    "bookSection": "chapter",
    "conferencePaper": "paper-conference",
    "report": "report",
    "thesis": "thesis",
}

_FIELD_MAP: dict[str, str] = {
    "title": "title",
    "publicationTitle": "container-title",
    "volume": "volume",
    "issue": "issue",
    "pages": "page",
    "DOI": "DOI",
    "ISSN": "ISSN",
    "ISBN": "ISBN",
    "abstractNote": "abstract",
    "url": "URL",
    "publisher": "publisher",
}


def _parse_date(date_str: str) -> dict:
    """Parse a date string into CSL-JSON date format.

    Handles ``2024``, ``2024-03``, ``2024-03-15``, and ``/``-separated
    variants.

    Args:
        date_str: Date string from Zotero.

    Returns:
        Dict with ``date-parts`` key.
    """
    if not date_str:
        return {"date-parts": [[]]}
    parts_str = re.split(r"[-/\s]", date_str)
    parts = [int(p) for p in parts_str if p.isdigit()]
    return {"date-parts": [parts]} if parts else {"date-parts": [[]]}


def zotero_to_csl_json(item_data: dict, user_id: str) -> dict:
    """Convert Zotero item data to CSL-JSON format.

    Maps Zotero fields to CSL-JSON fields for embedding in Word field codes.

    Args:
        item_data: Zotero item metadata dict.
        user_id: Zotero user ID for URI construction.

    Returns:
        CSL-JSON dict suitable for embedding in a Zotero field code.
    """
    item_key = item_data.get("key", "")
    item_type = item_data.get("itemType", "")

    csl: dict = {
        "type": _ITEM_TYPE_MAP.get(item_type, "article"),
        "id": int(hashlib.md5(item_key.encode()).hexdigest()[:8], 16),
        "_uris": [f"http://zotero.org/users/{user_id}/items/{item_key}"],
    }

    # Map simple fields
    for zotero_field, csl_field in _FIELD_MAP.items():
        value = item_data.get(zotero_field)
        if value:
            csl[csl_field] = value

    # Creators -- authors only
    creators = item_data.get("creators", [])
    authors = [
        {"family": c.get("lastName", ""), "given": c.get("firstName", "")}
        for c in creators
        if c.get("creatorType") == "author"
    ]
    if authors:
        csl["author"] = authors

    # Date
    date_str = item_data.get("date", "")
    if date_str:
        csl["issued"] = _parse_date(date_str)

    return csl


# ---------------------------------------------------------------------------
# 3. Field Code Builder
# ---------------------------------------------------------------------------


def _make_run() -> OxmlElement:
    """Create a new ``w:r`` element."""
    return OxmlElement("w:r")


def _make_fld_char(fld_char_type: str) -> OxmlElement:
    """Create a ``w:fldChar`` element with the given type.

    Args:
        fld_char_type: One of ``begin``, ``separate``, ``end``.

    Returns:
        A ``w:r`` element containing the fldChar.
    """
    run = _make_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), fld_char_type)
    run.append(fld_char)
    return run


def _make_instr_text(text: str) -> OxmlElement:
    """Create a ``w:r`` element containing ``w:instrText``.

    Args:
        text: The instruction text content.

    Returns:
        A ``w:r`` element with instrText child.
    """
    run = _make_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = text
    run.append(instr)
    return run


def _make_superscript_run(text: str) -> OxmlElement:
    """Create a ``w:r`` element with superscript text.

    Args:
        text: Display text to render as superscript.

    Returns:
        A ``w:r`` element with superscript formatting.
    """
    run = _make_run()
    rpr = OxmlElement("w:rPr")
    vert_align = OxmlElement("w:vertAlign")
    vert_align.set(qn("w:val"), "superscript")
    rpr.append(vert_align)
    run.append(rpr)
    t_elem = OxmlElement("w:t")
    t_elem.set(qn("xml:space"), "preserve")
    t_elem.text = text
    run.append(t_elem)
    return run


def add_citation_field(paragraph, citation_json: dict, display_text: str) -> None:
    """Insert a Zotero citation field code into a Word paragraph.

    Builds the XML structure:
      begin fldChar -> instrText -> separate fldChar -> display run -> end fldChar

    Args:
        paragraph: A python-docx Paragraph object.
        citation_json: The full citation JSON dict for the field code.
        display_text: Text to display (typically the Vancouver number).
    """
    p_elem = paragraph._element
    json_str = json.dumps(citation_json, ensure_ascii=False)
    instr_content = f"ADDIN ZOTERO_ITEM CSL_CITATION {json_str}"

    p_elem.append(_make_fld_char("begin"))
    p_elem.append(_make_instr_text(instr_content))
    p_elem.append(_make_fld_char("separate"))
    p_elem.append(_make_superscript_run(display_text))
    p_elem.append(_make_fld_char("end"))


def add_bibliography_field(paragraph) -> None:
    """Insert a Zotero bibliography field code into a Word paragraph.

    Args:
        paragraph: A python-docx Paragraph object.
    """
    p_elem = paragraph._element
    instr_content = 'ADDIN ZOTERO_BIBL {"uncited":[],"custom":[]} CSL_BIBLIOGRAPHY'

    p_elem.append(_make_fld_char("begin"))
    p_elem.append(_make_instr_text(instr_content))
    p_elem.append(_make_fld_char("separate"))
    p_elem.append(_make_fld_char("end"))


# ---------------------------------------------------------------------------
# 4. Document Assembler
# ---------------------------------------------------------------------------

# Regex for inline markdown formatting
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"\*(.+?)\*")


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text with basic markdown formatting (bold/italic) to a paragraph.

    Processes ``**bold**`` and ``*italic*`` markers. Text without markers
    is added as plain runs.

    Args:
        paragraph: A python-docx Paragraph object.
        text: Text that may contain markdown bold/italic markers.
    """
    # Split on bold markers first
    parts = _BOLD_RE.split(text)
    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 1:
            # Bold segment
            run = paragraph.add_run(part)
            run.bold = True
        else:
            # Check for italic within non-bold segments
            italic_parts = _ITALIC_RE.split(part)
            for j, ipart in enumerate(italic_parts):
                if not ipart:
                    continue
                if j % 2 == 1:
                    run = paragraph.add_run(ipart)
                    run.italic = True
                else:
                    paragraph.add_run(ipart)


def build_document(
    content: str,
    item_data: dict[str, dict],
    user_id: str,
    output_path: str,
) -> str:
    """Build a Word document with live Zotero citations.

    Args:
        content: Markdown text with ``[@KEY]`` citation markers.
        item_data: Dict mapping item keys to their Zotero metadata.
        user_id: Zotero user ID for URI construction.
        output_path: Where to save the .docx file.

    Returns:
        Absolute path of the saved file.
    """
    doc = Document()

    # Parse global citation numbering across the whole document
    _, key_to_number = parse_citations(content)

    # Split into paragraphs on blank lines
    raw_paragraphs = re.split(r"\n\n+", content.strip())

    for raw_para in raw_paragraphs:
        raw_para = raw_para.strip()
        if not raw_para:
            continue

        # Detect headings
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", raw_para)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            doc.add_heading(heading_text, level=level)
            continue

        # Regular paragraph -- parse citations within it
        # Use global key_to_number for consistent numbering across paragraphs
        blocks, _ = parse_citations(raw_para)
        for block in blocks:
            if block.kind == "citation":
                block.numbers = [key_to_number[k] for k in block.keys if k in key_to_number]

        paragraph = doc.add_paragraph()

        for block in blocks:
            if block.kind == "text":
                _add_formatted_text(paragraph, block.content)
            elif block.kind == "citation":
                # Build the citation field code
                citation_items = []
                for key in block.keys:
                    if key in item_data:
                        csl = zotero_to_csl_json(item_data[key], user_id)
                        uris = csl.pop("_uris", [])
                        citation_items.append(
                            {
                                "id": csl["id"],
                                "uris": uris,
                                "itemData": csl,
                            }
                        )

                # Vancouver display: e.g. "1" or "1,2"
                display = ",".join(str(n) for n in block.numbers)

                citation_json = {
                    "citationID": f"cite_{uuid.uuid4().hex[:8]}",
                    "properties": {
                        "formattedCitation": display,
                        "plainCitation": display,
                        "noteIndex": 0,
                    },
                    "citationItems": citation_items,
                    "schema": (
                        "https://github.com/citation-style-language"
                        "/schema/raw/master/csl-citation.json"
                    ),
                }
                add_citation_field(paragraph, citation_json, display)

    # Add References heading and bibliography
    doc.add_heading("References", level=1)
    bib_para = doc.add_paragraph()
    add_bibliography_field(bib_para)

    # Save
    output = Path(output_path).resolve()
    doc.save(str(output))
    return str(output)


# ---------------------------------------------------------------------------
# 5. In-place Citation Insertion (preserves existing document formatting)
# ---------------------------------------------------------------------------


def _paragraph_full_text(paragraph) -> str:
    """Extract full text from a paragraph including all runs.

    Args:
        paragraph: A python-docx Paragraph object.

    Returns:
        Concatenated text of all runs.
    """
    return "".join(run.text for run in paragraph.runs)


def _has_citation_markers(text: str) -> bool:
    """Check whether text contains [@KEY] citation markers.

    Args:
        text: Text to check.

    Returns:
        True if citation markers are found.
    """
    return bool(_CITATION_RE.search(text))


def _rebuild_paragraph_with_citations(
    paragraph,
    blocks: list[TextBlock],
    item_data: dict[str, dict],
    user_id: str,
) -> None:
    """Replace a paragraph's content with citation field codes in-place.

    Clears existing runs and rebuilds with text blocks and Zotero field codes.
    Preserves the paragraph's style (heading level, alignment, spacing, etc.).

    Args:
        paragraph: A python-docx Paragraph object to modify.
        blocks: Parsed TextBlocks from parse_citations.
        item_data: Dict mapping item keys to Zotero metadata.
        user_id: Zotero user ID for URI construction.
    """
    # Preserve paragraph style before clearing
    style = paragraph.style

    # Remove all existing runs from the paragraph XML
    p_elem = paragraph._element
    for child in list(p_elem):
        if child.tag == qn("w:r"):
            p_elem.remove(child)

    # Restore style
    paragraph.style = style

    for block in blocks:
        if block.kind == "text":
            _add_formatted_text(paragraph, block.content)
        elif block.kind == "citation":
            citation_items = []
            for key in block.keys:
                if key in item_data:
                    csl = zotero_to_csl_json(item_data[key], user_id)
                    uris = csl.pop("_uris", [])
                    citation_items.append(
                        {
                            "id": csl["id"],
                            "uris": uris,
                            "itemData": csl,
                        }
                    )

            display = ",".join(str(n) for n in block.numbers)

            citation_json = {
                "citationID": f"cite_{uuid.uuid4().hex[:8]}",
                "properties": {
                    "formattedCitation": display,
                    "plainCitation": display,
                    "noteIndex": 0,
                },
                "citationItems": citation_items,
                "schema": (
                    "https://github.com/citation-style-language/schema/raw/master/csl-citation.json"
                ),
            }
            add_citation_field(paragraph, citation_json, display)


def insert_citations(
    document_path: str,
    item_data: dict[str, dict],
    user_id: str,
    output_path: str | None = None,
) -> tuple[str, int]:
    """Insert Zotero citation field codes into an existing Word document.

    Opens the document, scans all paragraphs for [@KEY] markers, replaces
    them with live Zotero field codes, and appends a bibliography if one
    is not already present. All other document formatting (styles, headers,
    footers, images, tables, page layout) is preserved.

    Args:
        document_path: Path to the existing .docx file.
        item_data: Dict mapping item keys to their Zotero metadata.
        user_id: Zotero user ID for URI construction.
        output_path: Where to save. If None, overwrites the original.

    Returns:
        Tuple of (saved file path, number of citation markers replaced).
    """
    doc = Document(document_path)
    save_to = Path(output_path or document_path).resolve()

    # First pass: collect all citation keys across the entire document
    # for consistent Vancouver numbering
    all_text_parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = _paragraph_full_text(paragraph)
        if _has_citation_markers(text):
            all_text_parts.append(text)

    # Also scan tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = _paragraph_full_text(paragraph)
                    if _has_citation_markers(text):
                        all_text_parts.append(text)

    if not all_text_parts:
        doc.save(str(save_to))
        return str(save_to), 0

    # Build global numbering from concatenated text
    combined = "\n\n".join(all_text_parts)
    _, key_to_number = parse_citations(combined)
    citation_count = len(key_to_number)

    # Second pass: rebuild paragraphs that contain citation markers
    def _process_paragraph(paragraph) -> None:
        text = _paragraph_full_text(paragraph)
        if not _has_citation_markers(text):
            return
        blocks, _ = parse_citations(text)
        # Remap to global numbering
        for block in blocks:
            if block.kind == "citation":
                block.numbers = [key_to_number[k] for k in block.keys if k in key_to_number]
        _rebuild_paragraph_with_citations(paragraph, blocks, item_data, user_id)

    for paragraph in doc.paragraphs:
        _process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _process_paragraph(paragraph)

    # Add bibliography if not already present
    has_bibliography = False
    for paragraph in doc.paragraphs:
        if "ADDIN ZOTERO_BIBL" in paragraph._element.xml:
            has_bibliography = True
            break

    if not has_bibliography:
        doc.add_heading("References", level=1)
        bib_para = doc.add_paragraph()
        add_bibliography_field(bib_para)

    doc.save(str(save_to))
    return str(save_to), citation_count
