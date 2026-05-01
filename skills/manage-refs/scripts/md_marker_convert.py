#!/usr/bin/env python3
"""md_marker_convert.py — Convert Vancouver-style ``[N]`` / ``[N, M]`` markers in
markdown or .docx into pandoc-style ``[@KEY]`` / ``[@KEY1, @KEY2]`` citekeys
(or back), driven by an explicit N→key mapping.

Two directions:
  --to-keys     ``[N]``  → ``[@KEY]``   (default)
  --to-numbers  ``[@KEY]`` → ``[N]``    (round-trip / debug)

Inputs:
  --input  FILE        path to .md or .docx
  --output FILE        write converted file (extension must match input)
  --map    FILE        N↔key mapping (JSON ``{"1": "ABC123", ...}`` or 2-column CSV ``n,key``)
  --active-ns CSV      optional comma list of N's to convert; markers outside
                       this set are left untouched. Use to stage partial
                       conversion (e.g. sample-only build).

Numbers without a mapping (or outside ``--active-ns``) are left as plain
``[N]`` so a downstream Zotero refresh / hand-edit can finalize them.

Anti-Hallucination:
  - The mapping is the single source of truth. Never invents keys.
  - If a marker contains a number not in the map AND no ``--active-ns`` is
    given, it is left untouched and reported on stderr (exit 0). The caller
    must fix the map or accept the partial conversion.

Origin: generalized from the RFA-Adjunct ``build_v4_zotero_docx.py`` replacer
(2026-05-01), validated on a 21-reference manuscript.
"""

from __future__ import annotations

import argparse
import csv
import json
import re
import sys
from pathlib import Path

# [N], [N, M], [N,M,...], possibly with internal whitespace
_NUM_RE = re.compile(r"\[(\d+(?:\s*,\s*\d+)*)\]")
# Pandoc citekey group: [@KEY], [@K1, @K2], [@K1; @K2]
_KEY_RE = re.compile(r"\[(@[A-Za-z0-9_]+(?:\s*[,;]\s*@[A-Za-z0-9_]+)*)\]")


def load_map(path: Path) -> dict[int, str]:
    raw = path.read_text(encoding="utf-8").strip()
    if path.suffix.lower() == ".json" or raw.startswith("{"):
        data = json.loads(raw)
        return {int(k): str(v).strip() for k, v in data.items() if str(v).strip()}
    out: dict[int, str] = {}
    for row in csv.reader(raw.splitlines()):
        if not row or row[0].strip().lower() in {"n", "number", ""}:
            continue
        if len(row) < 2:
            continue
        out[int(row[0].strip())] = row[1].strip()
    return out


def parse_active(spec: str | None, full_keys: set[int]) -> set[int]:
    if not spec:
        return full_keys
    return {int(x.strip()) for x in spec.split(",") if x.strip()}


def make_num_to_key(n_to_key: dict[int, str], active: set[int]):
    unmapped: set[int] = set()

    def repl(m: re.Match) -> str:
        nums = [int(x.strip()) for x in m.group(1).split(",")]
        if not all(n in active for n in nums):
            return m.group(0)
        keys: list[str] = []
        for n in nums:
            k = n_to_key.get(n)
            if k is None:
                unmapped.add(n)
                return m.group(0)
            keys.append(f"@{k}")
        return "[" + ", ".join(keys) + "]"

    return repl, unmapped


def make_key_to_num(n_to_key: dict[int, str]):
    key_to_n = {v: n for n, v in n_to_key.items()}
    unknown: set[str] = set()

    def repl(m: re.Match) -> str:
        keys = [k.strip().lstrip("@") for k in re.split(r"[,;]", m.group(1))]
        nums: list[int] = []
        for k in keys:
            n = key_to_n.get(k)
            if n is None:
                unknown.add(k)
                return m.group(0)
            nums.append(n)
        return "[" + ", ".join(str(n) for n in nums) + "]"

    return repl, unknown


def transform_text(text: str, repl) -> str:
    # Auto-pick regex based on whether the text contains '@' citekeys.
    pattern = _KEY_RE if "[@" in text else _NUM_RE
    return pattern.sub(repl, text)


def convert_markdown(src: Path, dst: Path, repl) -> int:
    text = src.read_text(encoding="utf-8")
    pattern = _KEY_RE if "[@" in text else _NUM_RE
    n = sum(1 for _ in pattern.finditer(text))
    new = pattern.sub(repl, text)
    dst.write_text(new, encoding="utf-8")
    return n


def convert_docx(src: Path, dst: Path, repl, direction: str) -> int:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        sys.exit("ERROR: python-docx is required for .docx input. `pip install python-docx`")
    doc = Document(str(src))
    pattern = _KEY_RE if direction == "to-numbers" else _NUM_RE
    n = 0

    def process(p):
        nonlocal n
        for run in p.runs:
            if "[" not in run.text:
                continue
            n += sum(1 for _ in pattern.finditer(run.text))
            run.text = pattern.sub(repl, run.text)

    for p in doc.paragraphs:
        process(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process(p)
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    return n


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--input", required=True, type=Path)
    ap.add_argument("--output", required=True, type=Path)
    ap.add_argument("--map", required=True, type=Path)
    ap.add_argument("--active-ns", default=None,
                    help="Comma-separated N's to convert (default: all in map).")
    direction = ap.add_mutually_exclusive_group()
    direction.add_argument("--to-keys", action="store_true", default=True,
                           help="[N] → [@KEY] (default).")
    direction.add_argument("--to-numbers", action="store_true",
                           help="[@KEY] → [N] (round-trip).")
    args = ap.parse_args()

    if not args.input.exists():
        sys.exit(f"ERROR: input not found: {args.input}")
    if args.input.suffix != args.output.suffix:
        sys.exit("ERROR: --input and --output must share the same extension (.md or .docx).")

    n_to_key = load_map(args.map)
    if not n_to_key:
        sys.exit("ERROR: empty mapping.")

    if args.to_numbers:
        repl, unknown = make_key_to_num(n_to_key)
        direction = "to-numbers"
    else:
        active = parse_active(args.active_ns, set(n_to_key))
        repl, unknown = make_num_to_key(n_to_key, active)
        direction = "to-keys"

    if args.input.suffix == ".docx":
        seen = convert_docx(args.input, args.output, repl, direction)
    else:
        seen = convert_markdown(args.input, args.output, repl)

    print(f"[md_marker_convert] direction={direction} markers_seen={seen} → {args.output}",
          file=sys.stderr)
    if unknown:
        print(f"[md_marker_convert] WARNING: unmapped tokens left untouched: "
              f"{sorted(unknown)}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
